import sys
from docx import Document

def pularLinha(p, qtd): 
  run = p.add_run()
  for x in range(qtd):
    run.add_break()

def main():
  print('Iniciando')
  document = Document()
  font = document.styles['Normal'].font
  font.name = 'Arial'
  
  p = document.add_paragraph('')
  
  for row in sys.argv:
    if(row.__contains__('Pontos')):
      document.add_heading(row, 0)
      p = document.add_paragraph('')
      continue
    
    if(row.__contains__('Toque')):
      p.add_run(row).bold = True
      pularLinha(p, 1)
      continue
    elif(row.__contains__('Entidade')):
      p.add_run(row).bold = True
      pularLinha(p, 1)
      continue
    else:  
      p.add_run(row)
      pularLinha(p, 2)
      continue

  document.save('outputs/pontos.docx')

  print('Finalizado')

if __name__ == '__main__':
    main()